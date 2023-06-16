import sys
import docx
import PyPDF2
from PyQt5.QtWidgets import QApplication, QMainWindow, QMessageBox, QFileDialog
from giaodien import Ui_MainWindow
import CreateKey
import decode
import encode
import hashlib

class MainWindow:
    def __init__(self):
        self.main_win = QMainWindow()
        self.uic = Ui_MainWindow()
        self.uic.setupUi(self.main_win)
        self.uic.pushButtonmh1.clicked.connect(self.taokhoa)
        self.uic.pushButtonmh2.clicked.connect(self.mahoaRSA)
        self.uic.buttongiaimamh3.clicked.connect(self.giaimaRSA)
        self.uic.pushButton.clicked.connect(self.chonfile1)
        self.uic.buttongiaimamh5.clicked.connect(self.chonfile2)
        self.uic.pushButton_3.clicked.connect(self.saveFile1)
        self.uic.buttongiaimamh4.clicked.connect(self.saveFile2)
        self.uic.pushButton_2.clicked.connect(self.chuyenBanMa)

    def taokhoa(self):
        CreateKey.main()

        of = open("Data/PublicKey1.txt", "r", encoding="utf8")
        b = of.readline()
        n = of.readline()
        pub = "+" + str(b) + "+" + str(n)
        self.uic.publickeymh1.setText(pub)
        of.close()
        of = open("Data/PrivateKey1.txt", "r", encoding="utf8")
        a = of.readline()
        p = of.readline()
        q = of.readline()
        pri = "+" + str(a) + "+" + str(p) + "+" + str(q)
        self.uic.privatekeymh1.setText(pri)
        of.close()

    def mahoaRSA(self):
        banro = self.uic.textnhapmh2.toPlainText()
        of = open("Data/Plaintext.txt", "w", encoding="utf8")
        of.write(banro)
        of.close()
        c = encode.main()
        self.uic.texthienthimh2.setText(c)

    def giaimaRSA(self):
        banma = self.uic.textmh4.toPlainText()
        of = open("Data/CiphertextDecode.txt", "w", encoding="utf8")
        of.write(banma)
        of.close()
        c = decode.main()
        self.uic.textmh3.setText(c)

    def chonfile1(self):
        msg = QMessageBox()
        msg.setWindowTitle("Error")
        file, _ = QFileDialog.getOpenFileNames(None, "Chon file", "", "All Files (*)")

        if len(file) != 1:
            msg.setText("Chọn file bạn muốn ký")
            msg.exec_()
        else:
            a = file[0]
            duoi = a[a.rfind(".") + 1:]
            if duoi == "txt":
                of = open(file[0], "r", encoding="utf8")
                a = of.read()
                of.close()
                # print(a)
                self.uic.textnhapmh2.setText(a)
            elif duoi == "docx":
                banro = docx.Document(file[0])
                content = banro.paragraphs
                ban_ro = content[0].text
                for i in range(1, len(content)):
                    ban_ro += '\n'
                    ban_ro += content[i].text
                self.uic.textnhapmh2.setText(ban_ro)

            elif duoi == "pdf":
                banro = PyPDF2.PdfFileReader(file[0])
                self.uic.textnhapmh2.setText(banro.getPage(0).extractText())
    def chonfile2(self):
        msg = QMessageBox()
        msg.setWindowTitle("Error")
        file, _ = QFileDialog.getOpenFileNames(None, "Chon file", "", "All Files (*)")

        if len(file) != 1:
            msg.setText("Chọn file bạn muốn ký")
            msg.exec_()
        else:
            a = file[0]
            duoi = a[a.rfind(".") + 1:]
            if duoi == "txt":
                of = open(file[0], "r", encoding="utf8")
                a = of.read()
                of.close()
                # print(a)
                self.uic.textmh4.setText(a)
            elif duoi == "docx":
                banma = docx.Document(file[0])
                content = banma.paragraphs
                ban_ma = content[0].text
                for i in range(1, len(content)):
                    ban_ma += '\n'
                    ban_ma += content[i].text
                self.uic.textmh4.setText(ban_ma)

            elif duoi == "pdf":
                banma = PyPDF2.PdfFileReader(file[0])
                self.uic.textmh4.setText(banma.getPage(0).extractText())

    def chuyenBanMa(self):
        banma = self.uic.texthienthimh2.toPlainText()
        of = open("Data/Ciphertext.txt", "w", encoding="utf8")
        of.write(banma)
        of.close()
        c = encode.main()
        self.uic.texthienthimh2.clear()
        self.uic.textmh4.setText(c)

    def saveFile1(self):
        fileFormat, _ = QFileDialog.getSaveFileName(None, "Lưu file", "",
                                                    "All Files (*);;Text Files (*.txt);;Word Files (*.docx)")

        if fileFormat == "":
            return

        # Lấy nội dung hiện tại trong texthienthimh2
        text = self.uic.texthienthimh2.toPlainText()

        if fileFormat.endswith(".txt"):
            # Lưu văn bản thành file txt
            with open(fileFormat, "w", encoding="utf-8") as f:
                f.write(text)
        elif fileFormat.endswith(".docx"):
            # Lưu văn bản thành file Word
            doc = docx.Document()
            doc.add_paragraph(text)
            doc.save(fileFormat)

    def saveFile2(self):
        fileFormat, _ = QFileDialog.getSaveFileName(None, "Lưu file", "",
                                                    "All Files (*);;Text Files (*.txt);;Word Files (*.docx)")

        if fileFormat == "":
            return

        # Lấy nội dung hiện tại trong textmh3
        text = self.uic.textmh3.toPlainText()

        if fileFormat.endswith(".txt"):
            # Lưu văn bản thành file txt
            with open(fileFormat, "w", encoding="utf-8") as f:
                f.write(text)
        elif fileFormat.endswith(".docx"):
            # Lưu văn bản thành file Word
            doc = docx.Document()
            doc.add_paragraph(text)
            doc.save(fileFormat)

    def show(self):
        self.main_win.show()


if __name__ == "__main__":
    app = QApplication(sys.argv)
    main_win = MainWindow()
    main_win.show()
    sys.exit(app.exec())
